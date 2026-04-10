"""Resume DOCX renderer — clone the template and surgically replace bullets.

The template at ~/OneDrive/Desktop/talon-assistant/docs/Aaron_Lafferty_Resume_2026CRV1.docx
provides all formatting (fonts, margins, list styles, headers, education,
certifications). This renderer:

  1. Loads the template
  2. Replaces CAREER HIGHLIGHTS bullets with library picks
  3. Inserts CORE COMPETENCIES bullets from leadership_scope picks
  4. For each role, replaces bullets with library picks
  5. Replaces Talon tagline, bullets, and closing line
  6. Saves the new DOCX, then converts to PDF via docx2pdf

The Senior Analyst A&F entry and Education/Certifications stay verbatim.
Bullets are cloned from existing paragraphs so list numbering, spacing, and
font formatting are preserved automatically.
"""
from __future__ import annotations

import logging
import shutil
import tempfile
from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from core.resume_builder import ResumeLibrary, Selection

log = logging.getLogger(__name__)


DEFAULT_TEMPLATE_PATH = (
    Path.home() / "OneDrive" / "Desktop" / "talon-assistant" / "docs"
    / "Aaron_Lafferty_Resume_2026CRV1.docx"
)


# library slug -> company-header marker (matched as startswith on paragraph text)
ROLE_MARKERS: list[tuple[str, str]] = [
    ("amherst", "Amherst Group"),
    ("welldyne", "WellDyneRx"),
    ("cognizant", "Cognizant Technology Solutions"),
    ("abercrombie", "Abercrombie & Fitch"),     # first match = Manager role
    ("oarnet", "OARnet"),
]

CAREER_HIGHLIGHTS_HEADER = "CAREER HIGHLIGHTS"
CORE_COMPETENCIES_HEADER = "CORE COMPETENCIES"
TALON_HEADER = "AI DEVELOPMENT PROJECT"
EDUCATION_HEADER = "EDUCATION"


# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _is_bullet(p) -> bool:
    """A paragraph is a bullet if its pPr contains numPr."""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        return False
    return pPr.find(qn('w:numPr')) is not None


_XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'


def _split_bold_lead(text: str) -> tuple[str, str]:
    """Split a bullet into (bold_lead, rest) for the bold lead-in style.

    Tries strong punctuation, then strong prepositions, then weak prepositions,
    then commas, then a short last-resort split, then a 5-word fallback.
    Returns (text, '') if no reasonable split exists (very short bullets).
    """
    def best(breaks: tuple[str, ...], min_w: int, max_w: int):
        best_idx = None
        for brk in breaks:
            idx = text.find(brk)
            if idx <= 0:
                continue
            wc = len(text[:idx].split())
            if min_w <= wc <= max_w and (best_idx is None or idx < best_idx):
                best_idx = idx
        if best_idx is None:
            return None
        return text[:best_idx], text[best_idx:]

    # Tier 1: strong punctuation
    r = best((': ', '. '), 2, 13)
    if r:
        return r
    # Tier 2: strong prepositions
    r = best((
        ' by ', ' through ', ' across ', ' spanning ', ' using ',
        ' covering ', ' serving ', ' enabling ', ' satisfying ',
        ' protecting ', ' supporting ', ' against ', ' under ',
    ), 3, 10)
    if r:
        return r
    # Tier 3: weak prepositions, clause boundaries, participial modifiers
    r = best((
        ' for ', ' with ', ' from ', ' into ', ' including ',
        ' that ', ' which ',
        ' aligned ', ' based ', ' grounded ', ' focused ', ' built ',
        ' designed ', ' powered ',
    ), 3, 11)
    if r:
        return r
    # Tier 4: comma
    r = best((', ',), 3, 12)
    if r:
        return r
    # Tier 5: short last resort
    r = best((' of ', ' to ', ' in ', ' on ', ' at '), 1, 4)
    if r:
        return r
    # Final fallback: first 5 words
    words = text.split()
    if len(words) >= 5:
        bold = ' '.join(words[:5])
        return bold, text[len(bold):]
    return text, ''


def _make_t(parent_run, text: str):
    t = etree.SubElement(parent_run, qn('w:t'))
    t.text = text
    t.set(_XML_SPACE, 'preserve')
    return t


def _strip_bold_from_rpr(rPr) -> None:
    if rPr is None:
        return
    for tag in ('w:b', 'w:bCs'):
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)


def _ensure_bold_in_rpr(rPr) -> None:
    if rPr is None:
        return
    if rPr.find(qn('w:b')) is None:
        b = etree.SubElement(rPr, qn('w:b'))  # noqa: F841


def _set_paragraph_text(p_elem, text: str, *, allow_bold_lead: bool = True) -> None:
    """Replace the paragraph's runs with new content for `text`.

    If `allow_bold_lead` is True (default), splits the text into a bold
    lead-in run + a normal trailing run, preserving the first run's rPr
    formatting (font, color, size) on both runs. If no reasonable split
    is found, falls back to a single non-bold run.
    """
    runs = p_elem.findall(qn('w:r'))
    if not runs:
        r = etree.SubElement(p_elem, qn('w:r'))
        _make_t(r, text)
        return

    first = runs[0]
    rPr = first.find(qn('w:rPr'))
    # Drop everything inside the first run except its rPr
    for child in list(first):
        if child.tag != qn('w:rPr'):
            first.remove(child)
    # Drop sibling runs
    for r in runs[1:]:
        p_elem.remove(r)

    bold_lead, rest = ('', text)
    if allow_bold_lead:
        bold_lead, rest = _split_bold_lead(text)

    if bold_lead and rest:
        # Two-run path: first bold, second normal
        if rPr is None:
            rPr = etree.SubElement(first, qn('w:rPr'))
        _ensure_bold_in_rpr(rPr)
        # Make sure rPr stays first
        first.remove(rPr)
        first.insert(0, rPr)
        _make_t(first, bold_lead)

        second = etree.SubElement(p_elem, qn('w:r'))
        rPr2 = deepcopy(rPr)
        _strip_bold_from_rpr(rPr2)
        second.insert(0, rPr2)
        _make_t(second, rest)
    else:
        # Single-run path: strip bold so flat text doesn't render bolded
        _strip_bold_from_rpr(rPr)
        if rPr is not None:
            first.remove(rPr)
            first.insert(0, rPr)
        _make_t(first, text)


def _replace_block(template_para, count_old: int, new_texts: list[str]) -> None:
    """Replace `count_old` consecutive paragraphs starting at `template_para`
    with new clones of `template_para`, one per text in `new_texts`.
    """
    template_el = template_para._element
    parent = template_el.getparent()

    # Collect old siblings
    old_elements = []
    cur = template_el
    for _ in range(count_old):
        if cur is None:
            break
        old_elements.append(cur)
        cur = cur.getnext()

    # Insertion index
    idx = list(parent).index(template_el)

    # Build clones with new text
    for text in new_texts:
        clone = deepcopy(template_el)
        _set_paragraph_text(clone, text)
        parent.insert(idx, clone)
        idx += 1

    # Remove old
    for oe in old_elements:
        parent.remove(oe)


def _insert_block_after(anchor_para, template_para, new_texts: list[str]) -> None:
    """Insert a new block of paragraphs (clones of template_para) immediately
    after `anchor_para`, without removing anything.
    """
    anchor_el = anchor_para._element
    parent = anchor_el.getparent()
    idx = list(parent).index(anchor_el) + 1
    template_el = template_para._element
    for text in new_texts:
        clone = deepcopy(template_el)
        _set_paragraph_text(clone, text)
        parent.insert(idx, clone)
        idx += 1


# ─────────────────────────────────────────────────────────────────────────────
# Section locators
# ─────────────────────────────────────────────────────────────────────────────

def _find_paragraph_by_text(doc, marker: str, *, exact_upper: bool = False):
    for p in doc.paragraphs:
        txt = p.text.strip()
        if exact_upper:
            if txt.upper() == marker.upper():
                return p
        else:
            if txt.startswith(marker):
                return p
    return None


def _find_bullet_block_after(doc, anchor_para):
    """Walk forward from `anchor_para` to find the first bullet paragraph and
    count consecutive bullets. Returns (first_bullet_para, count).
    """
    paragraphs = doc.paragraphs
    anchor_el = anchor_para._element
    # Find anchor index in doc.paragraphs (object identity)
    start = None
    for i, p in enumerate(paragraphs):
        if p._element is anchor_el:
            start = i
            break
    if start is None:
        return None, 0

    # Walk forward to first bullet
    j = start + 1
    while j < len(paragraphs) and not _is_bullet(paragraphs[j]):
        # Stop early if we hit a section header (all caps + bold)
        if paragraphs[j].text.strip() and paragraphs[j].text.strip().isupper():
            return None, 0
        j += 1
    if j >= len(paragraphs):
        return None, 0

    first_bullet = paragraphs[j]
    count = 0
    k = j
    while k < len(paragraphs) and _is_bullet(paragraphs[k]):
        count += 1
        k += 1
    return first_bullet, count


# ─────────────────────────────────────────────────────────────────────────────
# Talon section helpers
# ─────────────────────────────────────────────────────────────────────────────

def _replace_italic_run_text(p_elem, text: str) -> bool:
    """Find the first italic run in p_elem and set its text to `text`.
    Strips any soft line breaks before it. Returns True on success.
    """
    runs = p_elem.findall(qn('w:r'))
    target = None
    for r in runs:
        rPr = r.find(qn('w:rPr'))
        if rPr is not None and rPr.find(qn('w:i')) is not None:
            target = r
            break
    if target is None:
        return False
    # Drop any <w:br> and <w:t> children, then set fresh <w:t>
    rPr = target.find(qn('w:rPr'))
    for child in list(target):
        if child.tag != qn('w:rPr'):
            target.remove(child)
    # Re-add a soft line break (so the blurb keeps its line under the title)
    etree.SubElement(target, qn('w:br'))
    t = etree.SubElement(target, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    if rPr is not None:
        target.remove(rPr)
        target.insert(0, rPr)
    return True


def _replace_full_paragraph_text(p_elem, text: str) -> None:
    """Used for the Talon italic closing line — keep it italic, no bold lead."""
    _set_paragraph_text(p_elem, text, allow_bold_lead=False)


# ─────────────────────────────────────────────────────────────────────────────
# Public renderer
# ─────────────────────────────────────────────────────────────────────────────

class TemplateRenderer:
    def __init__(self, template_path: Path | None = None) -> None:
        self.template_path = Path(template_path) if template_path else DEFAULT_TEMPLATE_PATH

    def render(
        self,
        library: ResumeLibrary,
        selection: Selection,
        out_path: Path,
    ) -> Path:
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")

        # OneDrive cloud-only files refuse Path.read_bytes() but shutil.copyfile
        # works. Stage a local copy first to avoid Document() permission errors.
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tf:
            staged = Path(tf.name)
        shutil.copyfile(self.template_path, staged)

        try:
            doc = Document(staged)

            self._replace_career_highlights(doc, library, selection)
            # Core Competencies is skipped: the template uses a table for this
            # section that we shouldn't duplicate. leadership_scope library
            # picks remain available for cover letter use.
            self._replace_role_bullets(doc, library, selection)
            self._replace_talon_section(doc, library, selection)

            out_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(out_path)
        finally:
            try:
                staged.unlink()
            except OSError:
                pass

        return out_path

    # ── Replacement steps ──────────────────────────────────────────────────

    def _replace_career_highlights(
        self, doc, library: ResumeLibrary, selection: Selection
    ) -> None:
        picks = sorted(selection.picks.get("career_highlights", []))
        section = library.get("career_highlights")
        if not section or not picks:
            return
        header = _find_paragraph_by_text(doc, CAREER_HIGHLIGHTS_HEADER, exact_upper=True)
        if header is None:
            log.warning("[ResumeDocx] CAREER HIGHLIGHTS header not found")
            return
        first_bullet, count = _find_bullet_block_after(doc, header)
        if first_bullet is None or count == 0:
            log.warning("[ResumeDocx] No bullets found under CAREER HIGHLIGHTS")
            return
        new_texts = [section.bullets[i] for i in picks if 0 <= i < len(section.bullets)]
        if new_texts:
            _replace_block(first_bullet, count, new_texts)
            log.info(f"[ResumeDocx] Career highlights: {count} -> {len(new_texts)}")

    def _insert_core_competencies(
        self, doc, library: ResumeLibrary, selection: Selection
    ) -> None:
        picks = sorted(selection.picks.get("leadership_scope", []))
        section = library.get("leadership_scope")
        if not section or not picks:
            return
        header = _find_paragraph_by_text(doc, CORE_COMPETENCIES_HEADER, exact_upper=True)
        if header is None:
            log.warning("[ResumeDocx] CORE COMPETENCIES header not found")
            return
        # Steal a bullet template from CAREER HIGHLIGHTS to clone formatting
        ch_header = _find_paragraph_by_text(doc, CAREER_HIGHLIGHTS_HEADER, exact_upper=True)
        bullet_template = None
        if ch_header is not None:
            bullet_template, _ = _find_bullet_block_after(doc, ch_header)
        if bullet_template is None:
            log.warning("[ResumeDocx] No bullet template available for Core Competencies")
            return
        # Check whether anything is already under the header (shouldn't be in
        # this template, but be defensive)
        existing_first, existing_count = _find_bullet_block_after(doc, header)
        new_texts = [section.bullets[i] for i in picks if 0 <= i < len(section.bullets)]
        if not new_texts:
            return
        if existing_first is not None and existing_count > 0:
            _replace_block(existing_first, existing_count, new_texts)
        else:
            _insert_block_after(header, bullet_template, new_texts)
        log.info(f"[ResumeDocx] Core competencies: inserted {len(new_texts)} bullets")

    def _replace_role_bullets(
        self, doc, library: ResumeLibrary, selection: Selection
    ) -> None:
        for slug, marker in ROLE_MARKERS:
            picks = sorted(selection.picks.get(slug, []))
            section = library.get(slug)
            if not section or not picks:
                continue
            anchor = _find_paragraph_by_text(doc, marker)
            if anchor is None:
                log.warning(f"[ResumeDocx] Role marker not found: {marker}")
                continue
            first_bullet, count = _find_bullet_block_after(doc, anchor)
            if first_bullet is None or count == 0:
                log.warning(f"[ResumeDocx] No bullets under role: {marker}")
                continue
            new_texts = [
                section.bullets[i] for i in picks if 0 <= i < len(section.bullets)
            ]
            if not new_texts:
                continue
            _replace_block(first_bullet, count, new_texts)
            log.info(f"[ResumeDocx] {slug}: {count} -> {len(new_texts)} bullets")

    def _replace_talon_section(
        self, doc, library: ResumeLibrary, selection: Selection
    ) -> None:
        talon = library.get("talon")
        if not talon:
            return
        header = _find_paragraph_by_text(doc, TALON_HEADER, exact_upper=True)
        if header is None:
            log.warning("[ResumeDocx] AI DEVELOPMENT PROJECT header not found")
            return

        # Walk forward: title-line para, then bullets, then italic closing,
        # stopping at EDUCATION.
        paragraphs = doc.paragraphs
        start = None
        for i, p in enumerate(paragraphs):
            if p._element is header._element:
                start = i
                break
        if start is None:
            return

        # Title/blurb is the first non-bullet paragraph after the header
        title_para = None
        for j in range(start + 1, len(paragraphs)):
            if paragraphs[j].text.strip().upper().startswith(EDUCATION_HEADER):
                break
            if not _is_bullet(paragraphs[j]) and paragraphs[j].text.strip():
                title_para = paragraphs[j]
                break

        # Update the italic blurb (tagline) in the title paragraph
        if title_para is not None and selection.talon_tagline is not None:
            taglines = talon.subsections.get("taglines", [])
            if 0 <= selection.talon_tagline < len(taglines):
                if not _replace_italic_run_text(
                    title_para._element, taglines[selection.talon_tagline]
                ):
                    log.warning("[ResumeDocx] Could not find italic run in Talon title para")

        # Find the bullet block (consecutive numPr paragraphs after the title)
        first_bullet, count = (None, 0)
        if title_para is not None:
            first_bullet, count = _find_bullet_block_after(doc, title_para)

        picks = sorted(selection.picks.get("talon", []))
        new_bullets = [talon.bullets[i] for i in picks if 0 <= i < len(talon.bullets)]
        if first_bullet is not None and count > 0 and new_bullets:
            _replace_block(first_bullet, count, new_bullets)
            log.info(f"[ResumeDocx] talon: {count} -> {len(new_bullets)} bullets")

        # Closing line: the next non-bullet paragraph after the bullets,
        # before EDUCATION. We re-walk doc.paragraphs because the prior step
        # mutated the tree.
        paragraphs = doc.paragraphs
        # Find Education header
        edu_idx = None
        for i, p in enumerate(paragraphs):
            if p.text.strip().upper().startswith(EDUCATION_HEADER):
                edu_idx = i
                break
        if edu_idx is None:
            return
        # Closing line is the last paragraph before edu that is not a bullet
        closing_para = None
        for k in range(edu_idx - 1, -1, -1):
            if _is_bullet(paragraphs[k]):
                break
            if paragraphs[k].text.strip():
                closing_para = paragraphs[k]
                break

        if closing_para is not None and selection.talon_closing is not None:
            closings = talon.subsections.get("closing", [])
            if 0 <= selection.talon_closing < len(closings):
                _replace_full_paragraph_text(
                    closing_para._element, closings[selection.talon_closing]
                )
                log.info("[ResumeDocx] talon closing line replaced")


# ─────────────────────────────────────────────────────────────────────────────
# Convenience entry point
# ─────────────────────────────────────────────────────────────────────────────

def render_resume_docx(
    library: ResumeLibrary,
    selection: Selection,
    out_path: Path,
    *,
    template_path: Path | None = None,
) -> Path:
    return TemplateRenderer(template_path).render(library, selection, out_path)


def convert_to_pdf(docx_path: Path) -> Path | None:
    """Convert a DOCX to PDF via docx2pdf (requires MS Word installed).

    Runs the conversion in a subprocess because docx2pdf uses COM
    automation which can raise a Windows fatal exception (0x800706be)
    that kills the entire Python process. Isolating it in a child
    process protects the main app from COM crashes.
    """
    import subprocess as _sp
    import sys

    pdf_path = docx_path.with_suffix(".pdf")

    # Run conversion in a child process to isolate COM crashes
    script = (
        "import sys; from docx2pdf import convert; "
        f"convert(r'{docx_path}', r'{pdf_path}')"
    )
    try:
        result = _sp.run(
            [sys.executable, "-c", script],
            capture_output=True, text=True, timeout=60,
        )
        if pdf_path.exists():
            return pdf_path
        if result.returncode != 0:
            log.warning(f"[ResumeDocx] PDF conversion subprocess failed "
                      f"(rc={result.returncode}): {result.stderr[:200]}")
    except _sp.TimeoutExpired:
        log.warning("[ResumeDocx] PDF conversion timed out (60s)")
    except Exception as e:
        log.warning(f"[ResumeDocx] PDF conversion failed: {e}")
    return None
