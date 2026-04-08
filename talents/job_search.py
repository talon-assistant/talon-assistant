"""JobSearchTalent -- automated job search across LinkedIn, Dice, and Built In.

Scrapes job listings from configured search URLs, identifies new postings
not yet in the job tracker, adds them automatically, and runs automated
resume-fit analysis via Claude CLI in the background.

Supported sites:
    - LinkedIn (requires one-time login via dedicated Chrome profile)
    - Dice.com (no auth required)
    - Built In / builtin.com (no auth required)

Designed to run on a schedule via the Talon scheduler — e.g.,
"schedule a job search every 2 hours"

Examples:
    "search for jobs"
    "run a job hunt"
    "check for new job listings"
    "add a search URL https://dice.com/jobs?q=..."
    "show my search URLs"
    "remove the second search URL"
    "job search login"
"""
from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import tempfile
import time
import threading
from datetime import date, datetime
from pathlib import Path
from typing import Any
from urllib.parse import urlparse, parse_qs, urlencode, urlunparse

from talents.base import BaseTalent

import logging
log = logging.getLogger(__name__)


def _data_dir() -> str:
    d = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data")
    os.makedirs(d, exist_ok=True)
    return d


def _detect_site(url: str) -> str:
    """Identify which job board a URL belongs to."""
    host = urlparse(url).netloc.lower()
    if "linkedin.com" in host:
        return "linkedin"
    if "dice.com" in host:
        return "dice"
    if "builtin.com" in host:
        return "builtin"
    return "unknown"


def _clean_search_url(url: str) -> str:
    """Strip ephemeral/session params from a search URL."""
    site = _detect_site(url)
    parsed = urlparse(url)
    params = parse_qs(parsed.query, keep_blank_values=True)

    if site == "linkedin":
        for key in ("currentJobId", "origin"):
            params.pop(key, None)

    # Flatten single-value lists for urlencode
    flat = {k: v[0] if len(v) == 1 else v for k, v in params.items()}
    return urlunparse(parsed._replace(query=urlencode(flat, doseq=True)))


class JobSearchTalent(BaseTalent):
    """Search LinkedIn, Dice, and Built In for job listings."""

    name = "job_search"
    description = (
        "Search for job listings on LinkedIn, Dice, and Built In; "
        "track new postings and hand matches to Cowork for fit analysis"
    )
    keywords = [
        "job search", "job searches", "job hunt", "job hunting",
        "hunt for jobs", "search for jobs", "find jobs", "find job",
        "check linkedin", "check dice", "new job listings",
        "search url", "search urls", "job listings",
        "todays job", "today's job", "do a job",
        "cover letter",
        "prepare materials", "tailor resume", "tailored resume",
        "customize resume", "customized resume",
        "prepare everything", "prep everything", "full application",
    ]
    examples = [
        "search for jobs",
        "run a job hunt",
        "do todays job searches",
        "check for new job listings",
        "find me some jobs",
        "show top jobs",
        "show best matches",
        "score my jobs",
        "run fit analysis",
        "add a search URL https://dice.com/jobs?q=...",
        "show my search URLs",
        "remove the first search URL",
        "job search login",
        "write a cover letter for #3",
        "prepare materials for #3",
        "tailor resume for #3",
        "prepare everything for #3",
    ]
    priority = 62
    required_packages = ["selenium"]

    def __init__(self) -> None:
        super().__init__()
        self._profile_dir = os.path.join(_data_dir(), "job_search_chrome_profile")
        self._config_file = os.path.join(_data_dir(), "job_search_config.json")
        self._search_config: dict[str, Any] = {"urls": [], "auto_cowork": True}
        self._load_search_config()
        log.info(f"[JobSearch] Config: {self._config_file} "
                 f"(exists={os.path.exists(self._config_file)}, "
                 f"urls={len(self._search_config.get('urls', []))})")

    # ── Persistent config ────────────────────────────────────────────────────

    def _load_search_config(self) -> None:
        if os.path.exists(self._config_file):
            try:
                with open(self._config_file, encoding="utf-8") as f:
                    self._search_config.update(json.load(f))
            except (json.JSONDecodeError, OSError) as e:
                log.warning(f"[JobSearch] Could not load config: {e}")

    def _save_search_config(self) -> None:
        os.makedirs(os.path.dirname(self._config_file), exist_ok=True)
        with open(self._config_file, "w", encoding="utf-8") as f:
            json.dump(self._search_config, f, indent=2)

    # ── GUI config schema ────────────────────────────────────────────────────

    def get_config_schema(self) -> dict:
        return {
            "fields": [
                {
                    "key": "auto_cowork",
                    "label": "Auto-run fit analysis on new listings via Claude CLI",
                    "type": "bool",
                    "default": True,
                },
            ]
        }

    # ── Main dispatch ────────────────────────────────────────────────────────

    def execute(self, command: str, context: dict) -> dict:
        cmd = command.lower()

        # URL management
        if "add" in cmd and ("url" in cmd or "search url" in cmd):
            return self._handle_add_url(command)
        if ("show" in cmd or "list" in cmd) and ("url" in cmd or "search" in cmd):
            return self._handle_list_urls()
        if ("remove" in cmd or "delete" in cmd) and ("url" in cmd or "search" in cmd):
            return self._handle_remove_url(command)

        # Login helper (LinkedIn only)
        if "login" in cmd:
            return self._handle_login()

        # "show top jobs" / "best matches" — handle directly
        if re.search(r'\b(top jobs|best match|top candidates|best jobs|'
                     r'top match|strongest|highest fit|best fit|ranked)\b',
                     cmd):
            return self._handle_top_candidates()

        # "score my jobs" / "run fit analysis" — score unscored existing jobs
        if re.search(r'\b(score|fit analysis|analyze|evaluate)\b', cmd):
            return self._handle_score_existing()

        # Full bundle: tailored resume + cover letter in one shot
        if ("prepare everything" in cmd
                or "prep everything" in cmd
                or "full application" in cmd):
            return self._handle_prepare_everything(command)

        # Tailored resume / full application materials via Claude CLI
        if ("prepare materials" in cmd
                or "tailor resume" in cmd
                or "tailored resume" in cmd
                or "customize resume" in cmd
                or "customized resume" in cmd):
            return self._handle_prepare_materials(command)

        # Cover letter generation via Claude CLI
        if "cover letter" in cmd:
            return self._handle_cover_letter(command)

        # Default: run the search
        return self._handle_search(context)

    # ── URL management ───────────────────────────────────────────────────────

    def _handle_add_url(self, command: str) -> dict:
        url_match = re.search(r'https?://[^\s]+', command)
        if not url_match:
            return _fail(
                "Include the full URL. Example: "
                "add search URL https://dice.com/jobs?q=security"
            )
        url = _clean_search_url(url_match.group())
        site = _detect_site(url)
        if site == "unknown":
            return _fail(
                "Unsupported site. Supported: LinkedIn, Dice, Built In."
            )
        self._search_config["urls"].append(url)
        self._save_search_config()
        n = len(self._search_config["urls"])
        return _ok(
            f"Added {site.title()} search URL ({n} total).",
            actions=[{"action": "add_search_url", "url": url, "site": site}],
        )

    def _handle_list_urls(self) -> dict:
        urls = self._search_config.get("urls", [])
        if not urls:
            return _ok("No search URLs configured. Say 'add a search URL' followed by the link.")
        lines = []
        for i, u in enumerate(urls, 1):
            site = _detect_site(u).title()
            display = u[:90] + "..." if len(u) > 90 else u
            lines.append(f"{i}. [{site}] {display}")
        return _ok("Search URLs:\n" + "\n".join(lines))

    def _handle_remove_url(self, command: str) -> dict:
        urls = self._search_config.get("urls", [])
        if not urls:
            return _fail("No search URLs to remove.")
        idx_match = re.search(r'(\d+)', command)
        if idx_match:
            idx = int(idx_match.group()) - 1
            if 0 <= idx < len(urls):
                urls.pop(idx)
                self._save_search_config()
                return _ok(f"Removed URL #{idx + 1}. {len(urls)} remaining.")
        return _fail(f"Specify which URL to remove (1\u2013{len(urls)}).")

    # ── Login (LinkedIn) ─────────────────────────────────────────────────────

    def _handle_login(self) -> dict:
        """Open Chrome with the dedicated profile so user can log into LinkedIn."""
        try:
            driver = self._create_driver_persistent(headless=False)
            driver.get("https://www.linkedin.com/login")
            return _ok(
                "Chrome opened to LinkedIn login. Sign in, then close the "
                "browser when done. Your session will be saved for future "
                "searches.",
                actions=[{"action": "login_prompt"}],
            )
        except Exception as e:
            return _fail(f"Could not open Chrome: {e}")

    # ── Selenium driver ──────────────────────────────────────────────────────

    def _build_driver(self, profile_dir: str, headless: bool, bypass_cache: bool):
        from selenium import webdriver
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.chrome.service import Service

        options = Options()
        options.add_argument(f"--user-data-dir={profile_dir}")
        options.add_argument("--disable-blink-features=AutomationControlled")
        options.add_argument("--window-size=1920,1080")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option("useAutomationExtension", False)

        if headless:
            options.add_argument("--headless=new")

        try:
            from webdriver_manager.chrome import ChromeDriverManager
            service = Service(ChromeDriverManager().install())
        except ImportError:
            service = Service()

        driver = webdriver.Chrome(service=service, options=options)
        driver.execute_cdp_cmd(
            "Page.addScriptToEvaluateOnNewDocument",
            {"source": "Object.defineProperty(navigator, 'webdriver', "
                       "{get: () => undefined})"},
        )
        if bypass_cache:
            # Kill the HTTP cache + service workers without nuking cookies.
            # Keeps the LinkedIn login intact while preventing stale search
            # results from a baked-in disk cache or stuck SW.
            try:
                driver.execute_cdp_cmd(
                    "Network.setCacheDisabled", {"cacheDisabled": True}
                )
                driver.execute_cdp_cmd(
                    "Network.setBypassServiceWorker", {"bypass": True}
                )
            except Exception as e:
                log.warning(f"[JobSearch] CDP cache-bypass failed: {e}")
        return driver

    def _create_driver_persistent(self, headless: bool = True):
        """Driver bound to the saved LinkedIn profile (keeps li_at cookie)."""
        return self._build_driver(self._profile_dir, headless, bypass_cache=True)

    def _create_driver_clean(self, headless: bool = True) -> tuple:
        """Driver with a throwaway profile. Returns (driver, temp_dir).

        Caller is responsible for ``shutil.rmtree(temp_dir, ignore_errors=True)``
        in a finally block. No login state, no cache pollution between runs.
        """
        temp_dir = tempfile.mkdtemp(prefix="talon_jobsearch_")
        try:
            driver = self._build_driver(temp_dir, headless, bypass_cache=False)
        except Exception:
            shutil.rmtree(temp_dir, ignore_errors=True)
            raise
        return driver, temp_dir

    # ── Score existing unscored jobs ────────────────────────────────────────

    def _handle_score_existing(self) -> dict:
        """Run fit analysis on tracker jobs that haven't been scored yet."""
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir

        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        if not os.path.exists(db_path):
            return _fail("No job tracker database found. Run a search first.")

        db = _DB(db_path)
        with db._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM applications WHERE archived = 0 "
                "AND status = 'new' AND fit_score = 0 "
                "ORDER BY id DESC"
            ).fetchall()
        unscored = [dict(r) for r in rows]

        if not unscored:
            return _ok("All tracked jobs already have fit scores.")

        # Convert to the format _run_fit_analysis expects
        jobs = []
        for app in unscored:
            jobs.append({
                "id": app["id"],
                "company": app["company"],
                "position": app["position"],
                "location": app.get("location", ""),
                "job_url": app.get("job_url", ""),
                "source": app.get("source", ""),
            })

        count = self._run_fit_analysis(jobs)
        return _ok(
            f"Scoring {count} unscored job(s) in the background. "
            f"Check back in a few minutes with 'show top jobs'."
        )

    # ── Shared helpers for materials / cover letter / prepare everything ──

    @staticmethod
    def _find_application(command: str, prefixes: tuple[str, ...]) -> tuple[dict | None, str | None]:
        """Resolve a tracker application from a command string.

        Returns (app, error_message). On success, error_message is None.
        Lookup order:
          1. Strict '#N' token (must be preceded by space/start; digits only)
          2. Company/position LIKE-search using the text after any known prefix
        If the LIKE-search returns more than one row, returns an error that
        lists the matches so the user can pick a specific ID.
        """
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir

        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        if not os.path.exists(db_path):
            return None, "No job tracker database found."
        db = _DB(db_path)

        # 1. Strict #N — leading '#' required so '3M' and '2026' don't trigger
        id_match = re.search(r'(?:^|\s)#(\d+)\b', command)
        if id_match:
            app = db.get_application(int(id_match.group(1)))
            if app:
                return app, None
            return None, f"No application found for #{id_match.group(1)}."

        # 2. Company-name search using a prefix anchor
        cmd_lower = command.lower()
        query = ""
        for prefix in prefixes:
            if prefix in cmd_lower:
                query = cmd_lower.split(prefix, 1)[-1].strip()
                break
        if not query:
            return None, None  # let caller supply the "which job?" message

        matches = db.search(query)
        if not matches:
            return None, f"No tracker matches for '{query}'."
        if len(matches) == 1:
            return matches[0], None
        # Multiple matches — force disambiguation
        lines = [f"Multiple matches for '{query}':"]
        for m in matches[:8]:
            lines.append(f"  #{m['id']} {m['company']} — {m['position']}")
        if len(matches) > 8:
            lines.append(f"  …and {len(matches) - 8} more")
        lines.append("Try again with '#<id>'.")
        return None, "\n".join(lines)

    def _get_or_fetch_jd(self, app: dict) -> str:
        """Return the JD text for an application.

        Prefers the persisted copy on the row; falls back to a live scrape,
        then persists it back to the DB so the next call is free.
        """
        stored = (app.get("job_description") or "").strip()
        if stored:
            return stored

        job_url = app.get("job_url", "")
        if not job_url:
            return ""

        jd = self._fetch_job_description(job_url)
        if jd:
            try:
                from talents.job_tracker import _DB, _data_dir as tracker_data_dir
                db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
                _DB(db_path).update_application(app["id"], job_description=jd)
                app["job_description"] = jd
            except Exception as e:
                log.warning(f"[JobSearch] Could not persist JD for #{app['id']}: {e}")
        return jd

    # ── Tailored resume (materials prep) via Claude CLI ─────────────────────

    def _handle_prepare_everything(self, command: str) -> dict:
        """Run both the tailored resume build and the cover letter in one
        shot. Returns a merged response listing every file produced.

        Command form: "prepare everything for #N" or
        "prepare everything for [company]".
        """
        # Resolve the app once so we can pre-fetch the JD and avoid two
        # back-to-back headless Chrome launches.
        app, err = self._find_application(command, (
            "prepare everything for", "prep everything for",
            "full application for",
        ))
        if err:
            return _fail(err)
        if app is None:
            return _fail(
                "Which job? Try: 'prepare everything for #3' "
                "or 'prepare everything for OEC'."
            )

        # Warm the JD cache on the row so both downstream handlers read from DB.
        self._get_or_fetch_jd(app)

        resume_res = self._handle_prepare_materials(command)
        letter_res = self._handle_cover_letter(command)

        resume_ok = bool(resume_res.get("success"))
        letter_ok = bool(letter_res.get("success"))
        resume_msg = resume_res.get("response", "") or ""
        letter_msg = letter_res.get("response", "") or ""

        # Total failure: bubble both errors up
        if not resume_ok and not letter_ok:
            return _fail(
                "Prepare everything failed on both steps.\n\n"
                f"Resume: {resume_msg}\n\nCover letter: {letter_msg}"
            )

        # Partial success: still return ok so files that were written aren't hidden
        if resume_ok and not letter_ok:
            return _ok(
                f"{resume_msg}\n\n"
                f"Cover letter step failed: {letter_msg}"
            )
        if letter_ok and not resume_ok:
            return _ok(
                f"Resume step failed: {resume_msg}\n\n"
                f"{letter_msg}"
            )

        # Full success: merge both payloads
        actions = []
        actions.extend(resume_res.get("actions_taken") or [])
        actions.extend(letter_res.get("actions_taken") or [])
        return {
            "success": True,
            "response": (
                "Full application package ready.\n\n"
                f"{resume_msg}\n\n"
                f"{letter_msg}"
            ),
            "actions_taken": actions,
            "spoken": "Full application package ready.",
        }

    def _handle_prepare_materials(self, command: str) -> dict:
        """Phase 1: build a tailored resume preview by selecting bullets
        from the pre-written library. No DOCX yet, markdown only.

        Command form: "prepare materials for #N" or "tailor resume for #N".
        """
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir
        from core.resume_builder import (
            DEFAULT_LIBRARY_PATH,
            ResumeLibrary,
            ResumeSelector,
            render_preview,
            render_selection_notes,
        )

        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        if not os.path.exists(db_path):
            return _fail("No job tracker database found.")
        db = _DB(db_path)

        # One-time idempotent cleanup of HTML entities in older rows
        try:
            n_cleaned = db.unescape_html_entities()
            if n_cleaned:
                log.info(f"[JobSearch] Unescaped HTML entities in {n_cleaned} rows")
        except Exception as e:
            log.warning(f"[JobSearch] Entity backfill failed: {e}")

        # Find application via shared helper
        app, err = self._find_application(command, (
            "prepare materials for", "tailor resume for",
            "tailored resume for", "customize resume for",
            "customized resume for", "prepare everything for",
            "prep everything for", "full application for",
        ))
        if err:
            return _fail(err)
        if app is None:
            return _fail(
                "Which job? Try: 'prepare materials for #3' "
                "or 'prepare materials for OEC'."
            )

        # Load library fresh
        lib_path = DEFAULT_LIBRARY_PATH
        if not lib_path.exists():
            return _fail(f"Resume library not found at {lib_path}")
        try:
            library = ResumeLibrary(lib_path).parse()
        except Exception as e:
            return _fail(f"Could not parse resume library: {e}")

        if not library.sections:
            return _fail("Resume library parsed but contained no sections.")

        # Pull JD (cached on the row if already fetched; scrape and persist otherwise)
        job_url = app.get("job_url", "")
        job_description = self._get_or_fetch_jd(app)
        if not job_description and app.get("notes"):
            job_description = app["notes"]

        # Run selector
        selector = ResumeSelector(library)
        try:
            selection = selector.pick(
                job_title=app.get("position", ""),
                company=app.get("company", ""),
                job_description=job_description or app.get("position", ""),
            )
        except Exception as e:
            log.error(f"[JobSearch] Resume selector failed: {e}")
            return _fail(f"Resume selector failed: {e}")

        if not selection.picks:
            return _fail(
                "Selector returned no bullet picks. "
                "Check the job description or the library format."
            )

        # Build output folder
        safe_company = _safe_slug(app.get('company', ''), max_len=60)
        safe_position = _safe_slug(app.get('position', ''), max_len=50)
        date_tag = date.today().strftime("%Y-%m-%d")
        out_root = Path.home() / "OneDrive" / "Documents" / "jobappmaterials"
        out_dir = out_root / f"{safe_company}_{safe_position}_{date_tag}"
        out_dir.mkdir(parents=True, exist_ok=True)

        preview_md = render_preview(
            library,
            selection,
            company=app.get("company", ""),
            job_title=app.get("position", ""),
            job_url=job_url,
        )
        notes_md = render_selection_notes(
            library,
            selection,
            company=app.get("company", ""),
            job_title=app.get("position", ""),
        )

        preview_path = out_dir / "resume_preview.md"
        notes_path = out_dir / "selection_notes.md"
        preview_path.write_text(preview_md, encoding="utf-8")
        notes_path.write_text(notes_md, encoding="utf-8")

        # Phase 2: render DOCX from the template, then convert to PDF
        from core.resume_docx import (
            DEFAULT_TEMPLATE_PATH,
            TemplateRenderer,
            convert_to_pdf,
        )

        saved_files = [str(preview_path), str(notes_path)]
        docx_path = out_dir / f"{safe_company}_{safe_position}_resume.docx"
        pdf_path: Path | None = None
        warnings: list[str] = []
        try:
            if not DEFAULT_TEMPLATE_PATH.exists():
                warnings.append(
                    f"Resume template missing: {DEFAULT_TEMPLATE_PATH} "
                    "(markdown files only)."
                )
                log.warning(
                    f"[JobSearch] Resume template missing: {DEFAULT_TEMPLATE_PATH}"
                )
            else:
                TemplateRenderer().render(library, selection, docx_path)
                saved_files.append(str(docx_path))
                try:
                    pdf_path = convert_to_pdf(docx_path)
                except Exception as pdf_err:
                    warnings.append(f"PDF conversion failed: {pdf_err}")
                    log.error(f"[JobSearch] PDF conversion failed: {pdf_err}")
                else:
                    if pdf_path:
                        saved_files.append(str(pdf_path))
                    else:
                        warnings.append("PDF conversion returned no file.")
        except Exception as e:
            warnings.append(f"DOCX render failed: {e}")
            log.error(f"[JobSearch] DOCX render failed: {e}")

        total_bullets = sum(len(v) for v in selection.picks.values())
        log.info(
            f"[JobSearch] Resume materials built for #{app['id']} "
            f"{app['company']} ({total_bullets} bullets)"
        )

        file_list = "\n".join(f"  {f}" for f in saved_files)
        formats = []
        if docx_path.exists():
            formats.append("DOCX")
        if pdf_path and pdf_path.exists():
            formats.append("PDF")
        formats_line = (
            f" ({' + '.join(formats)} ready for upload)" if formats else ""
        )

        warn_line = ""
        if warnings:
            warn_line = "\n\n\u26a0 " + " \u26a0 ".join(warnings)

        return _ok(
            f"Tailored materials for **{app['company']}** "
            f"({app['position']}) saved:\n\n"
            f"{file_list}\n\n"
            f"Picked {total_bullets} bullets across "
            f"{len(selection.picks)} sections{formats_line}."
            f"{warn_line}"
        )

    # ── Cover letter generation via Claude CLI ─────────────────────────────

    def _handle_cover_letter(self, command: str) -> dict:
        """Generate a cover letter for a tracked application."""
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir

        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        if not os.path.exists(db_path):
            return _fail("No job tracker database found.")

        db = _DB(db_path)

        app, err = self._find_application(command, (
            "write a cover letter for", "cover letter for",
            "write cover letter for", "generate a cover letter for",
            "create a cover letter for", "prepare everything for",
            "prep everything for", "full application for",
        ))
        if err:
            return _fail(err)
        if app is None:
            return _fail(
                "Which job? Try: 'write a cover letter for #3' "
                "or 'write a cover letter for OEC'."
            )

        # Read resume
        resume_path = Path.home() / "OneDrive" / "Documents" / "resume_master.md"
        try:
            resume_text = resume_path.read_text(encoding="utf-8")
        except Exception as e:
            return _fail(f"Cannot read resume: {e}")

        # Read style rules
        claude_md = Path.home() / ".claude" / "CLAUDE.md"
        style_rules = ""
        try:
            style_rules = claude_md.read_text(encoding="utf-8")
        except Exception:
            pass

        # Pull JD (cached on row if already fetched)
        job_url = app.get("job_url", "")
        job_description = self._get_or_fetch_jd(app)

        # Build prompt
        prompt_parts = [
            "TASK: Write a cover letter for the position below.",
            "",
            f"RESUME:\n{resume_text}",
            "",
            f"POSITION: {app['position']}",
            f"COMPANY: {app['company']}",
        ]
        if app.get("location"):
            prompt_parts.append(f"LOCATION: {app['location']}")
        if job_description:
            prompt_parts.append(f"\nJOB DESCRIPTION:\n{job_description}")
        if app.get("notes"):
            notes = app["notes"]
            if "Recommendation:" in notes:
                prompt_parts.append(f"\nFIT ANALYSIS: {notes}")

        prompt_parts.extend([
            "",
            f"STYLE RULES:\n{style_rules}" if style_rules else "",
            "",
            "INSTRUCTIONS:",
            "- Only promote matches and strengths. NEVER mention "
            "weaknesses, gaps, or missing qualifications.",
            "- No em dashes. Use commas, periods, or semicolons.",
            "- No tricolon / parallel triplets.",
            "- Plain language, direct executive tone.",
            "- Open with a specific hook tied to the company or role.",
            "- Pull specific metrics and accomplishments from the resume "
            "that match this role.",
            "- 3-4 paragraphs, under one page.",
            "- Close with confidence, not desperation.",
            "",
            "Output ONLY the cover letter text. No commentary, no "
            "preamble like 'Here is your cover letter', no markdown. "
            "Just the letter itself, starting with 'Dear'.",
        ])

        prompt = "\n".join(p for p in prompt_parts if p)

        # Call Claude CLI
        claude_bin = shutil.which("claude")
        if not claude_bin:
            return _fail("Claude CLI not found in PATH.")

        try:
            result = subprocess.run(
                [claude_bin, "-p", "--output-format", "text"],
                input=prompt,
                capture_output=True,
                text=True,
                encoding="utf-8",
                timeout=300,
                cwd=str(Path.home()),
            )

            if result.returncode != 0:
                log.error(f"[JobSearch] Cover letter failed: {result.stderr[:200]}")
                return _fail("Claude CLI failed to generate the cover letter.")

            letter = result.stdout.strip()
            if not letter or len(letter) < 50:
                return _fail("Claude returned an empty or too-short response.")

            # Sanity check
            _bad = ["webfetch", "permission", "tool call", "blocked by",
                    "approve the", "[your name]", "[your address]"]
            if any(s in letter.lower() for s in _bad):
                log.error(f"[JobSearch] Bad cover letter: {letter[:200]}")
                return _fail("Cover letter generation produced bad output.")

        except subprocess.TimeoutExpired:
            return _fail("Cover letter generation timed out.")

        # Save files
        output_dir = Path.home() / "OneDrive" / "Documents" / "Cover Letters"
        output_dir.mkdir(parents=True, exist_ok=True)

        import html as _html
        _c = _html.unescape(app.get('company', '') or '')
        _p = _html.unescape(app.get('position', '') or '')
        safe_company = re.sub(r'[^\w\s-]', '', _c).strip()
        safe_position = re.sub(r'[^\w\s-]', '', _p).strip()[:40]
        base_name = f"{safe_company} - {safe_position}"

        counter = 0
        suffix = ""
        while (output_dir / f"{base_name}{suffix}.docx").exists():
            counter += 1
            suffix = f" ({counter})"

        docx_path = output_dir / f"{base_name}{suffix}.docx"
        pdf_path = output_dir / f"{base_name}{suffix}.pdf"
        txt_path = output_dir / f"{base_name}{suffix}.txt"

        txt_path.write_text(letter, encoding="utf-8")
        saved = [str(txt_path)]
        warnings: list[str] = []

        try:
            self._write_cover_letter_docx(letter, docx_path, app)
            saved.append(str(docx_path))

            try:
                import docx2pdf
                docx2pdf.convert(str(docx_path), str(pdf_path))
                saved.append(str(pdf_path))
            except Exception as e:
                warnings.append(f"PDF conversion failed: {e}")
                log.warning(f"[JobSearch] PDF conversion failed: {e}")
        except Exception as e:
            warnings.append(f"DOCX generation failed: {e} (txt only)")
            log.warning(f"[JobSearch] DOCX generation failed: {e}")

        # Update tracker DB
        db.update_application(app["id"], cover_letter=1)
        log.info(f"[JobSearch] Cover letter saved for #{app['id']} {app['company']}")

        file_list = "\n".join(f"  {f}" for f in saved)
        formats_line = "DOCX and PDF ready for upload."
        if warnings:
            formats_line = "\u26a0 " + " \u26a0 ".join(warnings)

        return _ok(
            f"Cover letter for **{app['company']}** ({app['position']}) "
            f"saved to:\n\n{file_list}\n\n"
            f"{formats_line}"
        )

    def _fetch_job_description(self, job_url: str) -> str:
        """Scrape job description text from a URL using selenium.

        LinkedIn URLs use the persistent profile (auth required); everything
        else gets a throwaway profile so we don't pollute the login profile
        with cached ATS pages or contend for its lock.
        """
        try:
            from selenium.webdriver.common.by import By
        except Exception as e:
            log.warning(f"[JobSearch] Selenium unavailable: {e}")
            return ""

        is_linkedin = "linkedin.com" in (job_url or "").lower()
        temp_dir: str | None = None
        try:
            if is_linkedin:
                driver = self._create_driver_persistent(headless=True)
            else:
                driver, temp_dir = self._create_driver_clean(headless=True)
        except Exception as e:
            log.warning(f"[JobSearch] Could not start driver for JD: {e}")
            return ""

        try:
            driver.set_page_load_timeout(25)
            try:
                driver.get(job_url)
            except Exception as e:
                log.warning(f"[JobSearch] JD page load timed out: {e}")
            time.sleep(4)
            try:
                body = driver.find_element(By.TAG_NAME, "body").text
            except Exception as e:
                log.warning(f"[JobSearch] JD body read failed: {e}")
                return ""
            # Keep both ends of the JD: requirements often live mid/late,
            # qualifications at the top.
            if len(body) > 10000:
                body = body[:6000] + "\n[...truncated...]\n" + body[-4000:]
            log.info(f"[JobSearch] Fetched JD: {len(body)} chars")
            return body
        except Exception as e:
            log.warning(f"[JobSearch] Could not fetch JD: {e}")
            return ""
        finally:
            try:
                driver.quit()
            except Exception:
                pass
            if temp_dir:
                shutil.rmtree(temp_dir, ignore_errors=True)

    def scrape_and_add_from_url(self, job_url: str) -> dict:
        """Drop-in entry point for the Job Inbox dialog.

        Takes a raw URL the user pasted, scrapes the page for company +
        title + location + JD text, writes the row to the tracker DB, and
        kicks off background fit scoring. Returns a dict shaped like:
            {
                "success": bool,
                "error": str,          # present on failure
                "id": int,             # tracker row id (on success)
                "company": str,
                "position": str,
                "location": str,
                "job_url": str,
                "source": str,         # "LinkedIn" / "Dice" / "Built In" / "Manual"
                "duplicate": bool,     # True if the URL was already in the tracker
            }
        """
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir

        url = (job_url or "").strip()
        if not url or not url.startswith("http"):
            return {"success": False, "error": "Not a valid URL."}

        site = _detect_site(url)
        source_label = {
            "linkedin": "LinkedIn",
            "dice": "Dice",
            "builtin": "Built In",
        }.get(site, "Manual")

        # Dedup by bare URL before spinning up Chrome
        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        db = _DB(db_path)
        bare_url = url.split("?")[0]
        for existing in db.list_all(include_archived=True):
            if (existing.get("job_url") or "").split("?")[0] == bare_url:
                return {
                    "success": True,
                    "duplicate": True,
                    "id": existing["id"],
                    "company": existing.get("company", ""),
                    "position": existing.get("position", ""),
                    "location": existing.get("location", ""),
                    "job_url": existing.get("job_url", url),
                    "source": existing.get("source", source_label),
                }

        # Scrape the page
        parsed = self._scrape_posting_page(url)
        if not parsed or not parsed.get("position"):
            return {
                "success": False,
                "error": (
                    "Could not extract a job title from that URL. "
                    "Check the link or try opening it in Chrome first."
                ),
            }

        # Insert the row
        try:
            app_id = db.add_application(
                company=parsed.get("company") or "Unknown",
                position=parsed["position"],
                location=parsed.get("location", ""),
                source=source_label,
                status="new",
                date_found=date.today().isoformat(),
                job_url=url,
                job_description=parsed.get("job_description", ""),
            )
        except Exception as e:
            log.error(f"[JobSearch] Drop-in add failed: {e}")
            return {"success": False, "error": f"DB insert failed: {e}"}

        log.info(
            f"[JobSearch] Drop-in job added #{app_id}: "
            f"{parsed.get('company')} - {parsed['position']}"
        )

        # Kick off scoring in the background so the UI doesn't block
        self._run_fit_analysis([{
            "id": app_id,
            "company": parsed.get("company") or "Unknown",
            "position": parsed["position"],
            "location": parsed.get("location", ""),
            "job_url": url,
            "source": source_label,
        }])

        return {
            "success": True,
            "duplicate": False,
            "id": app_id,
            "company": parsed.get("company") or "Unknown",
            "position": parsed["position"],
            "location": parsed.get("location", ""),
            "job_url": url,
            "source": source_label,
        }

    def _scrape_posting_page(self, url: str) -> dict:
        """Scrape a single job posting URL for metadata + JD.

        Best-effort: pulls og:title / og:site_name from meta tags, falls
        back to <h1> and document.title, and grabs the visible body text
        for the JD. Works on most ATS pages and job boards without
        site-specific parsers.
        """
        from selenium.webdriver.common.by import By

        driver, temp_dir = self._create_driver_clean(headless=True)
        result: dict = {}
        try:
            driver.set_page_load_timeout(30)
            try:
                driver.get(url)
            except Exception as e:
                log.warning(f"[JobSearch] Drop-in page load warn: {e}")
            time.sleep(4)

            def _meta(name: str) -> str:
                try:
                    el = driver.find_element(
                        By.CSS_SELECTOR, f'meta[property="{name}"]'
                    )
                    return (el.get_attribute("content") or "").strip()
                except Exception:
                    return ""

            og_title = _meta("og:title")
            og_site = _meta("og:site_name")
            page_title = (driver.title or "").strip()

            # Position heuristic: prefer og:title, else first <h1>, else page title
            position = og_title
            if not position:
                try:
                    h1 = driver.find_element(By.TAG_NAME, "h1")
                    position = (h1.text or "").strip().split("\n")[0]
                except Exception:
                    pass
            if not position:
                position = page_title

            # Strip obvious boilerplate suffixes from the title
            for sep in (" | ", " - ", " – ", " — ", " at "):
                if sep in position and len(position) > 30:
                    # Usually "Title | Company" or "Title at Company"
                    left, right = position.split(sep, 1)
                    if not result.get("company") and 2 < len(right) < 80:
                        result["company"] = right.strip()
                    position = left.strip()
                    break

            result["position"] = position

            # Company heuristic: og:site_name, or common ATS company selectors
            if not result.get("company"):
                result["company"] = og_site
            if not result.get("company"):
                for sel in (
                    '[data-testid="inlineHeader-companyName"]',  # Greenhouse
                    '.company-name',
                    '[class*="company"]',
                    'a[href*="/company/"]',
                ):
                    try:
                        el = driver.find_element(By.CSS_SELECTOR, sel)
                        txt = (el.text or "").strip().split("\n")[0]
                        if txt and 1 < len(txt) < 80:
                            result["company"] = txt
                            break
                    except Exception:
                        pass

            # Location heuristic
            for sel in (
                '[data-testid*="location"]',
                '[class*="location"]',
                '.posting-categories .location',
            ):
                try:
                    el = driver.find_element(By.CSS_SELECTOR, sel)
                    txt = (el.text or "").strip().split("\n")[0]
                    if txt and len(txt) < 100:
                        result["location"] = txt
                        break
                except Exception:
                    pass

            # JD body
            try:
                body = driver.find_element(By.TAG_NAME, "body").text or ""
                if len(body) > 10000:
                    body = body[:6000] + "\n[...truncated...]\n" + body[-4000:]
                result["job_description"] = body
            except Exception:
                result["job_description"] = ""

        finally:
            try:
                driver.quit()
            finally:
                shutil.rmtree(temp_dir, ignore_errors=True)

        return result

    @staticmethod
    def _write_cover_letter_docx(letter_text: str, output_path: Path,
                                  app: dict) -> None:
        """Write a professionally formatted cover letter DOCX."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header.add_run("Aaron Lafferty")
        run.bold = True
        run.font.size = Pt(14)
        header.paragraph_format.space_after = Pt(2)

        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = contact.add_run(
            "aaron.lafferty@gmail.com | (614) 333-6612 | Irving, TX"
        )
        run.font.size = Pt(10)
        contact.paragraph_format.space_after = Pt(12)

        date_para = doc.add_paragraph()
        date_para.add_run(date.today().strftime("%B %d, %Y"))
        date_para.paragraph_format.space_after = Pt(12)

        paragraphs = [p.strip() for p in letter_text.split("\n\n") if p.strip()]
        for para_text in paragraphs:
            lower = para_text.lower()
            if lower.startswith("aaron lafferty"):
                continue
            clean_text = para_text.replace("\n", " ")
            p = doc.add_paragraph(clean_text)
            p.paragraph_format.space_after = Pt(8)

        doc.save(str(output_path))

    # ── Top candidates (queries tracker DB directly) ───────────────────────

    def _handle_top_candidates(self) -> dict:
        """Show new jobs ranked by fit score from the tracker DB."""
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir

        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        if not os.path.exists(db_path):
            return _fail("No job tracker database found. Run a search first.")

        db = _DB(db_path)
        apps = db.list_top_candidates(limit=15)

        if not apps:
            return _ok(
                "No scored candidates yet. Fit scores are added "
                "automatically after a job search. Try 'search for jobs' "
                "first, then check back in a few minutes."
            )

        lines = [f"**Top candidates** ({len(apps)}):\n"]
        for app in apps:
            loc = f" ({app['location']})" if app.get("location") else ""
            rec = ""
            notes = app.get("notes", "")
            if "Recommendation:" in notes:
                rec_part = notes.split("Recommendation:")[-1].strip()
                rec = f" [{rec_part}]"
            url = app.get("job_url", "")
            url_part = f" | {url}" if url else ""
            lines.append(
                f"  #{app['id']} [{app['fit_score']}%] "
                f"**{app['company']}** -- {app['position']}{loc}{rec}"
                f"{url_part}"
            )

        lines.append(
            "\nSay 'write a cover letter for #ID' or "
            "'write a cover letter for [company]' to generate one."
        )

        return _ok("\n".join(lines))

    # ── Search orchestrator ──────────────────────────────────────────────────

    def _handle_search(self, context: dict) -> dict:
        urls = self._search_config.get("urls", [])
        if not urls:
            return _fail(
                "No search URLs configured. Say 'add a search URL' "
                "followed by your search link."
            )

        all_jobs: list[dict] = []
        errors: list[str] = []
        sites_searched: set[str] = set()

        for url in urls:
            site = _detect_site(url)
            sites_searched.add(site)
            try:
                scraper = {
                    "linkedin": self._scrape_linkedin,
                    "dice":     self._scrape_dice,
                    "builtin":  self._scrape_builtin,
                }.get(site)
                if scraper:
                    jobs = scraper(url)
                    all_jobs.extend(jobs)
                else:
                    log.warning(f"[JobSearch] Unknown site for URL: {url[:60]}")
            except Exception as e:
                log.error(f"[JobSearch] Scrape failed for {url[:60]}: {e}")
                errors.append(str(e))

        if not all_jobs and errors:
            err_text = " ".join(errors).lower()
            if "login" in err_text or "authwall" in err_text:
                return _fail(
                    "LinkedIn requires login. Say 'job search login' to "
                    "open the browser and sign in."
                )
            return _fail(f"Search failed: {errors[0]}")

        sites_label = ", ".join(s.title() for s in sorted(sites_searched))
        if not all_jobs:
            return _ok(
                f"Searched {sites_label} ({len(urls)} URL(s)) \u2014 "
                "no listings found.",
                actions=[{"action": "search", "found": 0, "new": 0}],
            )

        # Dedup against existing tracker DB
        new_jobs = self._dedup_jobs(all_jobs)

        if not new_jobs:
            return _ok(
                f"Searched {sites_label}, found {len(all_jobs)} "
                f"listing(s) \u2014 all already tracked.",
                actions=[{"action": "search", "found": len(all_jobs), "new": 0}],
            )

        # Add new jobs to tracker DB
        added = self._add_to_tracker(new_jobs)

        # Automated fit analysis via Claude CLI
        fit_count = 0
        if self._search_config.get("auto_cowork") and added:
            fit_count = self._run_fit_analysis(added)

        # Build response
        lines = [f"Found {len(all_jobs)} listings across {sites_label}, "
                 f"{len(added)} new:"]
        for job in added[:10]:
            src = job.get("source", "")
            loc = f" ({job['location']})" if job.get("location") else ""
            lines.append(
                f"  \u2022 [{src}] {job.get('company', '?')} \u2014 "
                f"{job['position']}{loc}"
            )
        if len(added) > 10:
            lines.append(f"  \u2026 and {len(added) - 10} more")
        if fit_count:
            lines.append(
                f"\nFit analysis running in background for {fit_count} listing(s)."
            )

        notify = context.get("notify")
        if notify:
            notify("Job Search", f"{len(added)} new listing(s) found")

        return _ok(
            "\n".join(lines),
            actions=[{
                "action": "search",
                "found": len(all_jobs),
                "new": len(added),
                "fit_analysis": fit_count,
                "sites": list(sites_searched),
            }],
        )

    # ═══════════════════════════════════════════════════════════════════════════
    #  Site-specific scrapers — based on live DOM inspection (April 2026)
    # ═══════════════════════════════════════════════════════════════════════════

    # ── LinkedIn ─────────────────────────────────────────────────────────────
    #
    # Verified structure: div[data-job-id] cards inside
    # .scaffold-layout__list-container.  Each card contains:
    #   - a[href*="/jobs/view/"] for title + URL
    #   - .job-card-container__company-name for company
    #   - .job-card-container__metadata-item for location
    #
    # Both /jobs/search/ (classic) and /jobs/search-results/ (AI/semantic
    # search) render a card list in a left rail. The semantic-search page
    # wraps cards in <li data-occludable-job-id> instead of
    # <div data-job-id>, and scrolls a sub-panel instead of the page body,
    # so we accept multiple card selectors and walk the scroll container.

    _LINKEDIN_CARD_SELECTORS = (
        "div[data-job-id]",
        "li[data-occludable-job-id]",
        "div.job-card-job-posting-card-wrapper",
        "li.scaffold-layout__list-item",
    )
    _LINKEDIN_CARD_UNION = ", ".join(_LINKEDIN_CARD_SELECTORS)

    _LINKEDIN_SCROLL_CONTAINERS = (
        ".scaffold-layout__list",
        ".jobs-search-results-list",
        "div[data-results-list-top-scroll-sentinel] ~ ul",
    )

    def _scrape_linkedin(self, url: str) -> list[dict]:
        """Scrape job listings from a LinkedIn jobs search page."""
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC

        driver = self._create_driver_persistent(headless=True)
        jobs: list[dict] = []

        try:
            driver.get(url)
            time.sleep(5)

            if "/login" in driver.current_url or "authwall" in driver.current_url:
                raise RuntimeError(
                    "LinkedIn requires login \u2014 say 'job search login'"
                )

            # Wait for job cards or job links
            try:
                WebDriverWait(driver, 12).until(
                    EC.presence_of_element_located((
                        By.CSS_SELECTOR,
                        f'{self._LINKEDIN_CARD_UNION}, a[href*="/jobs/view/"]',
                    ))
                )
            except Exception:
                pass

            time.sleep(2)

            # ── Lazy-load loop: LinkedIn renders ~25 cards per
            # intersection-observer fire and only loads the next batch when
            # something near the bottom comes into view. The page body is
            # NOT the scroll container — the result list is a sub-panel
            # with its own overflow. So we (a) scrollTop the sub-panel,
            # (b) scrollIntoView the last card as a backup, and (c) cap
            # iterations so a stuck scroll can't hang us.
            scroll_js = (
                "const sels = arguments[0];"
                "for (const s of sels) {"
                "  const el = document.querySelector(s);"
                "  if (el && el.scrollHeight > el.clientHeight) {"
                "    el.scrollTop = el.scrollHeight;"
                "    return true;"
                "  }"
                "}"
                "return false;"
            )
            last_count = 0
            stable_passes = 0
            for attempt in range(20):
                cards = driver.find_elements(
                    By.CSS_SELECTOR, self._LINKEDIN_CARD_UNION
                )
                scrolled = False
                try:
                    scrolled = bool(driver.execute_script(
                        scroll_js, list(self._LINKEDIN_SCROLL_CONTAINERS)
                    ))
                except Exception:
                    pass
                if not scrolled and cards:
                    try:
                        driver.execute_script(
                            "arguments[0].scrollIntoView({block: 'end'});",
                            cards[-1],
                        )
                    except Exception:
                        pass
                time.sleep(1.5)
                new_count = len(driver.find_elements(
                    By.CSS_SELECTOR, self._LINKEDIN_CARD_UNION
                ))
                if new_count == last_count:
                    stable_passes += 1
                    if stable_passes >= 2:
                        break
                else:
                    stable_passes = 0
                last_count = new_count
            log.info(
                f"[JobSearch] LinkedIn scroll loop: {last_count} cards "
                f"after {attempt + 1} passes"
            )

            # Strategy 1: card-based (standard search page)
            cards = driver.find_elements(
                By.CSS_SELECTOR, self._LINKEDIN_CARD_UNION
            )
            if cards:
                seen: set[str] = set()
                for card in cards:
                    try:
                        job = self._parse_linkedin_card(card)
                        if job and job.get("position"):
                            uid = job.get("job_url", job["position"])
                            if uid not in seen:
                                seen.add(uid)
                                jobs.append(job)
                    except Exception as e:
                        log.debug(f"[JobSearch] LinkedIn card parse: {e}")

            # Strategy 2: link-based fallback (semantic search or alt layout).
            # Only accept rows where we can resolve a company from a nearby
            # ancestor element — otherwise we'd stuff the tracker with
            # half-empty "Unknown" rows that break dedup and scoring.
            if not jobs:
                seen_urls: set[str] = set()
                for link in driver.find_elements(
                    By.CSS_SELECTOR, 'a[href*="/jobs/view/"]'
                ):
                    text = link.text.strip()
                    if text:
                        text = text.split("\n")[0].strip()
                    href = (link.get_attribute("href") or "").split("?")[0]
                    if not text or len(text) < 4 or href in seen_urls:
                        continue

                    # Walk ancestors looking for a sibling that exposes
                    # the company name.
                    company = ""
                    location = ""
                    anc = link
                    for _ in range(6):
                        try:
                            anc = anc.find_element(By.XPATH, "..")
                        except Exception:
                            break
                        for sel in (
                            ".job-card-container__company-name",
                            ".artdeco-entity-lockup__subtitle",
                            ".job-card-container__primary-description",
                        ):
                            try:
                                el = anc.find_element(By.CSS_SELECTOR, sel)
                                txt = el.text.strip()
                                if txt and len(txt) > 1:
                                    company = txt
                                    break
                            except Exception:
                                pass
                        if company:
                            # Try to grab a location sibling too
                            for sel in (
                                ".job-card-container__metadata-item",
                                ".artdeco-entity-lockup__caption",
                            ):
                                try:
                                    el = anc.find_element(By.CSS_SELECTOR, sel)
                                    txt = el.text.strip()
                                    if txt:
                                        location = txt
                                        break
                                except Exception:
                                    pass
                            break

                    if not company:
                        log.debug(
                            f"[JobSearch] LinkedIn Strategy 2 skipped "
                            f"(no company): {text[:60]}"
                        )
                        continue

                    seen_urls.add(href)
                    jobs.append({
                        "source": "LinkedIn",
                        "date_found": date.today().isoformat(),
                        "position": text,
                        "company": company,
                        "location": location,
                        "job_url": href,
                    })

            # Strategy 3: CDP shadow-DOM piercing (SDUI / semantic search).
            # The /jobs/search-results/ endpoint mounts cards into closed
            # shadow roots, which Selenium's find_elements + page_source
            # can't see. CDP DOM.getDocument with pierce=True walks the
            # full composed tree including closed shadow roots.
            if not jobs:
                try:
                    cdp_jobs = self._scrape_linkedin_via_cdp(driver)
                    if cdp_jobs:
                        log.info(
                            f"[JobSearch] LinkedIn CDP fallback: "
                            f"{len(cdp_jobs)} listings"
                        )
                        jobs.extend(cdp_jobs)
                except Exception as e:
                    log.warning(f"[JobSearch] CDP fallback failed: {e}")

            log.info(f"[JobSearch] LinkedIn: {len(jobs)} listings")

            # If we got nothing, dump what headless Chrome actually saw so
            # we can debug selector / authwall / A/B-variant issues without
            # having to guess from a blind log line.
            if not jobs:
                try:
                    debug_dir = os.path.join(_data_dir(), "debug")
                    os.makedirs(debug_dir, exist_ok=True)
                    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    png_path = os.path.join(
                        debug_dir, f"linkedin_empty_{stamp}.png"
                    )
                    html_path = os.path.join(
                        debug_dir, f"linkedin_empty_{stamp}.html"
                    )
                    driver.save_screenshot(png_path)
                    with open(html_path, "w", encoding="utf-8") as f:
                        f.write(driver.page_source or "")
                    log.warning(
                        f"[JobSearch] LinkedIn returned 0 listings; "
                        f"dumped {png_path} and {html_path} "
                        f"(current_url={driver.current_url})"
                    )
                except Exception as e:
                    log.warning(f"[JobSearch] Debug dump failed: {e}")
        finally:
            driver.quit()

        return jobs

    @staticmethod
    def _parse_linkedin_card(card) -> dict | None:
        """Parse a LinkedIn div[data-job-id] card element."""
        from selenium.webdriver.common.by import By

        job: dict[str, str] = {
            "source": "LinkedIn",
            "date_found": date.today().isoformat(),
        }

        job_id = (
            card.get_attribute("data-job-id")
            or card.get_attribute("data-occludable-job-id")
        )
        if job_id:
            job["job_url"] = f"https://www.linkedin.com/jobs/view/{job_id}"

        # Title from the job link
        for link in card.find_elements(
            By.CSS_SELECTOR, 'a[href*="/jobs/view/"]'
        ):
            text = link.text.strip()
            # LinkedIn sometimes duplicates text (e.g. "Title\nTitle with verification")
            # Take first line only
            if text:
                text = text.split("\n")[0].strip()
            if text and 3 < len(text) < 200:
                job["position"] = text
                if not job.get("job_url"):
                    job["job_url"] = (link.get_attribute("href") or "").split("?")[0]
                break

        if not job.get("position"):
            return None

        # Company
        for sel in (".job-card-container__company-name",
                    ".artdeco-entity-lockup__subtitle",
                    ".job-card-container__primary-description"):
            for elem in card.find_elements(By.CSS_SELECTOR, sel):
                text = elem.text.strip()
                if text and len(text) > 1:
                    job["company"] = text
                    break
            if job.get("company"):
                break

        # Location
        for sel in (".job-card-container__metadata-item",
                    ".artdeco-entity-lockup__caption"):
            for elem in card.find_elements(By.CSS_SELECTOR, sel):
                text = elem.text.strip()
                if text and ("remote" in text.lower() or "," in text
                             or len(text) > 5):
                    job["location"] = text
                    break
            if job.get("location"):
                break

        return job

    # ── CDP shadow-DOM piercing for SDUI / semantic-search pages ─────────
    #
    # LinkedIn's /jobs/search-results/ AI-search endpoint mounts each card
    # into a closed shadow root via their Server-Driven UI framework.
    # Selenium's find_elements + page_source can't see closed shadow DOM,
    # but Chrome DevTools Protocol can: DOM.getDocument(pierce=True)
    # returns the full composed tree including every shadow root.
    #
    # The strategy:
    #   1. Get the pierced DOM tree as JSON.
    #   2. Build a parent map keyed by id() of each node dict.
    #   3. Find every <a href*="/jobs/view/"> anchor.
    #   4. For each anchor, walk up to ~10 ancestors collecting text
    #      until we have enough lines to parse out company + location.
    #   5. Skip "similar jobs" sub-sections by deduping bare URLs.

    @staticmethod
    def _cdp_attrs(node: dict) -> dict:
        """Convert CDP's flat attribute list ['k','v','k','v'] to a dict."""
        flat = node.get("attributes") or []
        return dict(zip(flat[::2], flat[1::2]))

    @classmethod
    def _cdp_walk(cls, node: dict, fn) -> None:
        """Depth-first walk over a CDP DOM tree, calling fn(node) per node.

        Pierces children, shadowRoots, contentDocument (iframes), and
        templateContent so closed shadow roots are visited too.
        """
        fn(node)
        for child in node.get("children") or ():
            cls._cdp_walk(child, fn)
        for shadow in node.get("shadowRoots") or ():
            cls._cdp_walk(shadow, fn)
        if node.get("contentDocument"):
            cls._cdp_walk(node["contentDocument"], fn)
        if node.get("templateContent"):
            cls._cdp_walk(node["templateContent"], fn)

    @classmethod
    def _cdp_collect_text(cls, node: dict, max_depth: int = 8) -> str:
        """Concatenate visible text under a node, separated by newlines.

        Splits at element boundaries so the output approximates the
        visual line layout of the card.
        """
        parts: list[str] = []

        def collect(n: dict, depth: int) -> None:
            if depth > max_depth:
                return
            node_type = n.get("nodeType")
            # nodeType 3 = TEXT_NODE
            if node_type == 3:
                v = (n.get("nodeValue") or "").strip()
                if v:
                    parts.append(v)
                return
            # Skip script/style/svg subtrees — pure noise
            name = (n.get("nodeName") or "").lower()
            if name in ("script", "style", "svg", "noscript"):
                return
            for c in n.get("children") or ():
                collect(c, depth + 1)
            for s in n.get("shadowRoots") or ():
                collect(s, depth + 1)

        collect(node, 0)
        return "\n".join(parts)

    @classmethod
    def _cdp_build_parent_map(cls, root: dict) -> dict:
        """Build {id(node_dict): parent_node_dict} for upward walking."""
        parents: dict = {}

        def visit(n: dict, parent: dict | None) -> None:
            parents[id(n)] = parent
            for c in n.get("children") or ():
                visit(c, n)
            for s in n.get("shadowRoots") or ():
                visit(s, n)
            if n.get("contentDocument"):
                visit(n["contentDocument"], n)

        visit(root, None)
        return parents

    # Lines that show up on every card and aren't useful for company/location
    _LI_NOISE_LINES = {
        "easy apply", "promoted", "saved", "be an early applicant",
        "viewed", "applied", "verified", "with verification",
        "promoted by hirer", "·", "•", "new", "actively reviewing",
        "responses managed off linkedin", "promoted job",
    }

    @classmethod
    def _clean_card_lines(cls, raw_text: str) -> list[str]:
        """Split a card's text dump into meaningful lines."""
        out: list[str] = []
        for ln in raw_text.split("\n"):
            ln = ln.strip()
            if not ln:
                continue
            low = ln.lower()
            if low in cls._LI_NOISE_LINES:
                continue
            if low.startswith("be an early") or low.startswith("posted "):
                continue
            if low.endswith(" ago") and len(ln) < 25:
                continue
            # Dedupe consecutive duplicates (LinkedIn often double-emits
            # the title for accessibility)
            if out and out[-1] == ln:
                continue
            out.append(ln)
        return out

    def _scrape_linkedin_via_cdp(self, driver) -> list[dict]:
        """Fallback scraper that walks the pierced DOM via CDP.

        Works on the SDUI /jobs/search-results/ page where cards live
        in closed shadow roots invisible to Selenium's normal API.
        """
        try:
            doc = driver.execute_cdp_cmd(
                "DOM.getDocument", {"depth": -1, "pierce": True}
            )
        except Exception as e:
            log.warning(f"[JobSearch] CDP getDocument failed: {e}")
            return []

        root = doc.get("root") or {}
        if not root:
            return []

        parents = self._cdp_build_parent_map(root)

        # Collect every anchor whose href looks like a job posting
        anchors: list[dict] = []
        seen_hrefs: set[str] = set()

        def visit(node: dict) -> None:
            if (node.get("nodeName") or "").lower() != "a":
                return
            attrs = self._cdp_attrs(node)
            href = attrs.get("href", "")
            if "/jobs/view/" not in href:
                return
            bare = href.split("?")[0]
            if bare in seen_hrefs:
                return
            seen_hrefs.add(bare)
            if not bare.startswith("http"):
                bare = "https://www.linkedin.com" + bare
            anchor_text = self._cdp_collect_text(node, max_depth=4)
            anchors.append({
                "node": node,
                "href": bare,
                "anchor_text": anchor_text,
            })

        self._cdp_walk(root, visit)

        log.info(
            f"[JobSearch] CDP found {len(anchors)} unique job anchors "
            f"in pierced DOM"
        )

        # Diagnostic: when we find suspiciously few job anchors, dump a
        # summary of what IS in the pierced tree so we can see whether
        # the cards use a different element type / URL pattern entirely.
        if len(anchors) < 3:
            try:
                self._dump_cdp_tree_summary(root)
            except Exception as e:
                log.warning(f"[JobSearch] CDP tree dump failed: {e}")

        jobs: list[dict] = []
        for entry in anchors:
            node = entry["node"]
            anchor_text = entry["anchor_text"]
            href = entry["href"]

            # Title: first non-empty line of the anchor's own text
            title_lines = self._clean_card_lines(anchor_text)
            title = title_lines[0] if title_lines else ""
            if not title or len(title) < 4 or len(title) > 200:
                continue

            # Walk up ancestors collecting card-level text until we get
            # enough lines to look like a real card row.
            company = ""
            location = ""
            ancestor = parents.get(id(node))
            for _ in range(10):
                if ancestor is None:
                    break
                lines = self._clean_card_lines(
                    self._cdp_collect_text(ancestor, max_depth=10)
                )
                if len(lines) >= 3:
                    # Find where the title sits in the line list
                    title_idx = -1
                    for i, ln in enumerate(lines):
                        if ln == title or title.lower() in ln.lower():
                            title_idx = i
                            break

                    # Company is usually the line right after the title
                    if 0 <= title_idx < len(lines) - 1:
                        candidate = lines[title_idx + 1]
                        if 1 < len(candidate) < 100:
                            company = candidate

                    # Location: first remaining line that mentions
                    # "remote" or contains a comma (City, ST format)
                    search_from = max(title_idx + 1, 0)
                    for ln in lines[search_from:]:
                        if ln == company:
                            continue
                        low = ln.lower()
                        if ("remote" in low or "," in ln) and len(ln) < 100:
                            location = ln
                            break

                    if company:
                        break
                ancestor = parents.get(id(ancestor))

            if not company:
                continue

            jobs.append({
                "source": "LinkedIn",
                "date_found": date.today().isoformat(),
                "position": title,
                "company": company,
                "location": location,
                "job_url": href,
            })

        return jobs

    @classmethod
    def _dump_cdp_tree_summary(cls, root: dict) -> None:
        """Write a CDP tree summary to data/debug for diagnostic use.

        Records:
          - tag frequency table (so we can spot non-anchor card elements)
          - all unique href values (so we can spot alternate URL patterns)
          - all unique role / data-testid / componentkey attribute values
          - frame / iframe contentDocument presence
        """
        from collections import Counter

        tags: Counter = Counter()
        hrefs: set[str] = set()
        roles: Counter = Counter()
        testids: Counter = Counter()
        component_keys: set[str] = set()
        sdui_components: set[str] = set()
        iframe_count = 0
        shadow_count = 0
        text_with_job: set[str] = set()

        def visit(node: dict) -> None:
            nonlocal iframe_count, shadow_count
            name = (node.get("nodeName") or "").lower()
            if name and not name.startswith("#"):
                tags[name] += 1
            attrs = cls._cdp_attrs(node)
            href = attrs.get("href")
            if href:
                hrefs.add(href[:200])
            role = attrs.get("role")
            if role:
                roles[role] += 1
            tid = attrs.get("data-testid")
            if tid:
                testids[tid] += 1
            ck = attrs.get("componentkey")
            if ck:
                component_keys.add(ck)
            sd = attrs.get("data-sdui-component")
            if sd:
                sdui_components.add(sd)
            if name == "iframe":
                iframe_count += 1
            if node.get("shadowRoots"):
                shadow_count += len(node["shadowRoots"])
            # Capture short text nodes that mention "job" (to find
            # rail markers without anchors)
            if node.get("nodeType") == 3:
                v = (node.get("nodeValue") or "").strip()
                if v and "job" in v.lower() and len(v) < 80:
                    text_with_job.add(v)

        cls._cdp_walk(root, visit)

        debug_dir = os.path.join(_data_dir(), "debug")
        os.makedirs(debug_dir, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = os.path.join(debug_dir, f"linkedin_cdp_summary_{stamp}.txt")

        lines = []
        lines.append(f"=== CDP pierced DOM summary @ {stamp} ===")
        lines.append(f"iframe nodes : {iframe_count}")
        lines.append(f"shadow roots : {shadow_count}")
        lines.append("")
        lines.append("--- top 30 tag names ---")
        for tag, count in tags.most_common(30):
            lines.append(f"  {count:5d}  {tag}")
        lines.append("")
        lines.append(f"--- {len(hrefs)} unique href values ---")
        for h in sorted(hrefs)[:80]:
            lines.append(f"  {h}")
        if len(hrefs) > 80:
            lines.append(f"  ... and {len(hrefs) - 80} more")
        lines.append("")
        lines.append(f"--- {len(roles)} unique role values ---")
        for role, count in roles.most_common(20):
            lines.append(f"  {count:5d}  role={role}")
        lines.append("")
        lines.append(f"--- {len(testids)} unique data-testid values ---")
        for tid, count in testids.most_common(40):
            lines.append(f"  {count:5d}  {tid}")
        lines.append("")
        lines.append(f"--- {len(component_keys)} componentkey values ---")
        for ck in sorted(component_keys)[:50]:
            lines.append(f"  {ck}")
        lines.append("")
        lines.append(f"--- {len(sdui_components)} sdui-component values ---")
        for sd in sorted(sdui_components)[:50]:
            lines.append(f"  {sd}")
        lines.append("")
        lines.append(f"--- text nodes mentioning 'job' ({len(text_with_job)}) ---")
        for t in sorted(text_with_job)[:60]:
            lines.append(f"  {t}")

        # Targeted: for the first 3 "(Verified job)" text nodes, walk
        # their ancestor chain and dump tag + attributes for each level
        # so we can see exactly where the card container is and which
        # attribute holds the job ID.
        verified_nodes: list[dict] = []

        def find_verified(node: dict) -> None:
            if node.get("nodeType") == 3:
                v = (node.get("nodeValue") or "").strip()
                if v.endswith("(Verified job)") and len(verified_nodes) < 3:
                    verified_nodes.append(node)

        cls._cdp_walk(root, find_verified)
        parents = cls._cdp_build_parent_map(root)

        if verified_nodes:
            lines.append("")
            lines.append("=== ancestor chains for first 3 job text nodes ===")
            for vn in verified_nodes:
                title_preview = (vn.get("nodeValue") or "")[:60]
                lines.append("")
                lines.append(f"--- text: {title_preview} ---")
                ancestor = parents.get(id(vn))
                level = 0
                while ancestor is not None and level < 12:
                    name = (ancestor.get("nodeName") or "?").lower()
                    attrs = cls._cdp_attrs(ancestor)
                    attr_strs = []
                    for k, v in attrs.items():
                        vs = str(v)
                        if len(vs) > 80:
                            vs = vs[:77] + "..."
                        attr_strs.append(f'{k}="{vs}"')
                    attr_blob = " ".join(attr_strs) if attr_strs else "(no attrs)"
                    lines.append(f"  L{level:2d} <{name}> {attr_blob}")
                    ancestor = parents.get(id(ancestor))
                    level += 1

        # Also: search for any 10+ digit numbers in any attribute value
        # anywhere in the tree (job IDs are typically 10 digits).
        import re as _re
        id_pat = _re.compile(r"\b\d{10}\b")
        found_ids: dict[str, set[str]] = {}

        def find_ids(node: dict) -> None:
            attrs = cls._cdp_attrs(node)
            for k, v in attrs.items():
                vs = str(v)
                for m in id_pat.findall(vs):
                    found_ids.setdefault(k, set()).add(m)

        cls._cdp_walk(root, find_ids)

        lines.append("")
        lines.append("=== 10-digit numbers in attributes (likely job IDs) ===")
        for k in sorted(found_ids.keys()):
            sample = sorted(found_ids[k])[:8]
            lines.append(
                f"  attr={k}  ({len(found_ids[k])} unique)  "
                f"sample={sample}"
            )

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        log.warning(f"[JobSearch] CDP tree summary dumped to {path}")

    # ── Dice ─────────────────────────────────────────────────────────────────
    #
    # Verified structure (April 2026): Dice is a JS SPA. Job cards are <div>
    # elements with Tailwind classes (rounded-lg border bg-surface-primary).
    # Each card contains a[href*="/job-detail/"] links. The card's full text
    # splits into lines like:
    #   [Company, "Easy Apply"/"Apply Now", Title, Location, dot, "Today"]
    #
    # The most reliable extraction: find all unique /job-detail/ URLs, then
    # walk up to the card-level parent and parse its text lines.

    def _scrape_dice(self, url: str) -> list[dict]:
        """Scrape job listings from Dice.com search results.

        Walks pages 1..N by setting ?page=K on the URL. Stops when a page
        adds zero new jobs or after the hard cap, whichever comes first.
        """
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC

        driver, temp_dir = self._create_driver_clean(headless=True)
        jobs: list[dict] = []
        seen_urls: set[str] = set()
        max_pages = 5

        try:
            for page_num in range(1, max_pages + 1):
                page_url = self._dice_url_with_page(url, page_num)
                driver.get(page_url)
                time.sleep(5)

                try:
                    WebDriverWait(driver, 15).until(
                        EC.presence_of_element_located((
                            By.CSS_SELECTOR,
                            'a[href*="/job-detail/"]',
                        ))
                    )
                except Exception:
                    pass
                time.sleep(3)

                links = driver.find_elements(
                    By.CSS_SELECTOR, 'a[href*="/job-detail/"]'
                )
                page_added = 0
                for link in links:
                    href = (link.get_attribute("href") or "").split("?")[0]
                    if not href or href in seen_urls:
                        continue

                    # Walk up to find the card container
                    card_el = link
                    for _ in range(8):
                        card_el = card_el.find_element(By.XPATH, "..")
                        card_text = card_el.text or ""
                        lines = [ln.strip() for ln in card_text.split("\n")
                                 if ln.strip()]
                        if len(lines) >= 3:
                            break
                    else:
                        continue

                    job = self._parse_dice_card_text(lines, href)
                    if job:
                        seen_urls.add(href)
                        jobs.append(job)
                        page_added += 1

                log.info(
                    f"[JobSearch] Dice page {page_num}: +{page_added} "
                    f"({len(jobs)} total)"
                )
                if page_added == 0:
                    break

            log.info(f"[JobSearch] Dice: {len(jobs)} listings")
        finally:
            try:
                driver.quit()
            finally:
                shutil.rmtree(temp_dir, ignore_errors=True)

        return jobs

    @staticmethod
    def _dice_url_with_page(url: str, page_num: int) -> str:
        """Return the Dice search URL with ?page=N set (replacing any prior)."""
        parsed = urlparse(url)
        params = parse_qs(parsed.query, keep_blank_values=True)
        params["page"] = [str(page_num)]
        flat = {k: v[0] if len(v) == 1 else v for k, v in params.items()}
        return urlunparse(parsed._replace(query=urlencode(flat, doseq=True)))

    @staticmethod
    def _parse_dice_card_text(lines: list[str], job_url: str) -> dict | None:
        """Parse a Dice job card from its text lines.

        Typical line order:
            [Company, "Easy Apply"/"Apply Now", Title, Location, dot, "Today"/"Yesterday"]
        But order can vary, so we use heuristics.
        """
        _SKIP = {"easy apply", "apply now", "today", "yesterday",
                 "\u00b7", "posted", "new"}

        # Filter noise lines
        meaningful = []
        for ln in lines:
            if ln.lower() in _SKIP or len(ln) <= 1:
                continue
            if ln.lower().startswith("posted "):
                continue
            meaningful.append(ln)

        if len(meaningful) < 2:
            return None

        # Heuristic: company is typically first, title is the longest
        # meaningful line, location contains comma or "Remote"
        company = meaningful[0]
        title = ""
        location = ""

        for ln in meaningful[1:]:
            if not title and len(ln) > 5:
                title = ln
            elif not location and ("," in ln or "remote" in ln.lower()):
                location = ln

        # If title looks like a location, swap
        if title and ("," in title and not location):
            # Check if next meaningful line is longer (more likely a title)
            remaining = [ln for ln in meaningful[1:] if ln != title]
            for r in remaining:
                if len(r) > len(title):
                    location = title
                    title = r
                    break

        if not title:
            return None

        return {
            "source": "Dice",
            "date_found": date.today().isoformat(),
            "position": title,
            "company": company,
            "location": location,
            "job_url": job_url,
        }

    # ── Built In ─────────────────────────────────────────────────────────────
    #
    # Verified structure (April 2026): Job links use a[href*="/job/"] pattern
    # with paths like /job/title-slug/12345.  The page is server-rendered so
    # links are available without waiting for JS.  Simple link extraction is
    # the most reliable approach.

    def _scrape_builtin(self, url: str) -> list[dict]:
        """Scrape job listings from builtin.com search results.

        Scrolls in a plateau loop until link count stops growing or the
        hard cap is hit. Built In server-renders the first batch and
        lazy-loads additional cards on scroll.
        """
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC

        driver, temp_dir = self._create_driver_clean(headless=True)
        jobs: list[dict] = []
        seen_urls: set[str] = set()

        def _collect_links() -> int:
            added = 0
            for link in driver.find_elements(
                By.CSS_SELECTOR, 'a[href*="/job/"]'
            ):
                text = link.text.strip()
                href = (link.get_attribute("href") or "").split("?")[0]
                if not text or len(text) < 4 or len(text) > 200:
                    continue
                if href in seen_urls:
                    continue
                if not href.startswith("http"):
                    href = "https://builtin.com" + href
                seen_urls.add(href)
                jobs.append({
                    "source": "Built In",
                    "date_found": date.today().isoformat(),
                    "position": text,
                    "job_url": href,
                })
                added += 1
            return added

        try:
            driver.get(url)
            time.sleep(4)

            try:
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((
                        By.CSS_SELECTOR, 'a[href*="/job/"]',
                    ))
                )
            except Exception:
                pass
            time.sleep(2)

            _collect_links()

            # Plateau scroll loop
            stable_passes = 0
            for attempt in range(6):
                try:
                    driver.execute_script(
                        "window.scrollTo(0, document.body.scrollHeight)"
                    )
                except Exception:
                    pass
                time.sleep(2.5)
                added = _collect_links()
                if added == 0:
                    stable_passes += 1
                    if stable_passes >= 2:
                        break
                else:
                    stable_passes = 0
            log.info(
                f"[JobSearch] Built In scroll loop: {len(jobs)} listings "
                f"after {attempt + 1} passes"
            )
        finally:
            try:
                driver.quit()
            finally:
                shutil.rmtree(temp_dir, ignore_errors=True)

        return jobs

    # ── Dedup against tracker DB ─────────────────────────────────────────────

    @staticmethod
    def _dedup_jobs(jobs: list[dict]) -> list[dict]:
        """Filter out jobs already present in the job tracker database."""
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir
        from talents.job_tracker import _normalize_company

        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        if not os.path.exists(db_path):
            return list(jobs)

        db = _DB(db_path)
        existing = db.list_all()

        existing_keys: set[tuple[str, str]] = set()
        existing_urls: set[str] = set()
        for app in existing:
            existing_keys.add((
                _normalize_company(app["company"]),
                app["position"].lower().strip(),
            ))
            if app.get("job_url"):
                existing_urls.add(app["job_url"].split("?")[0])

        new: list[dict] = []
        for job in jobs:
            url_base = (job.get("job_url") or "").split("?")[0]
            if url_base and url_base in existing_urls:
                continue
            key = (
                _normalize_company(job.get("company", "")),
                job.get("position", "").lower().strip(),
            )
            if key in existing_keys:
                continue
            new.append(job)
            existing_keys.add(key)
            if url_base:
                existing_urls.add(url_base)

        return new

    # ── Add to tracker DB ────────────────────────────────────────────────────

    @staticmethod
    def _add_to_tracker(jobs: list[dict]) -> list[dict]:
        """Insert new jobs into the job_tracker database."""
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir

        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        db = _DB(db_path)

        added: list[dict] = []
        for job in jobs:
            try:
                app_id = db.add_application(
                    company=job.get("company", "Unknown"),
                    position=job.get("position", ""),
                    location=job.get("location", ""),
                    source=job.get("source", ""),
                    status="new",
                    date_found=job.get("date_found", date.today().isoformat()),
                    job_url=job.get("job_url", ""),
                    salary_range=job.get("salary_range", ""),
                )
                job["id"] = app_id
                added.append(job)
            except Exception as e:
                log.error(f"[JobSearch] Failed to add {job.get('company')}: {e}")

        if added:
            log.info(f"[JobSearch] Added {len(added)} new applications to tracker")
        return added

    # ── Automated fit analysis via Claude CLI ──────────────────────────────

    def _run_fit_analysis(self, jobs: list[dict]) -> int:
        """Run fit analysis in a background thread using claude -p.

        Calls Claude CLI to score each job against the user's resume,
        then writes fit_score, notes, and recommendation back to the
        tracker database.  Runs in a daemon thread so it doesn't block
        the user response.
        """
        if not jobs:
            return 0

        count = len(jobs)
        thread = threading.Thread(
            target=self._fit_analysis_worker,
            args=(list(jobs),),
            daemon=True,
        )
        thread.start()
        log.info(f"[JobSearch] Fit analysis started in background ({count} jobs)")
        return count

    def _fit_analysis_worker(self, jobs: list[dict]) -> None:
        """Background worker: call claude -p for fit analysis, store results.

        Processes in batches of 5 to avoid output limits and URL fetch
        timeouts from overwhelming a single call.
        """
        batch_size = 5
        for i in range(0, len(jobs), batch_size):
            batch = jobs[i:i + batch_size]
            try:
                self._score_batch(batch)
            except Exception as e:
                log.error(f"[JobSearch] Batch {i // batch_size + 1} failed: {e}")

    def _score_batch(self, jobs: list[dict]) -> None:
        """Score a single batch of jobs via claude -p.

        For each job, scrapes the JD (if not already on the row) and
        persists it back to the DB so downstream steps (prepare materials,
        cover letter) can reuse it without another Chrome launch.
        """
        resume_path = Path.home() / "OneDrive" / "Documents" / "resume_master.md"

        # Read resume content directly so Claude doesn't need file access
        try:
            resume_text = resume_path.read_text(encoding="utf-8")
        except Exception as e:
            log.error(f"[JobSearch] Cannot read resume: {e}")
            return

        # Load the full rows so we have access to any already-stored JD
        # (jobs passed in from a fresh scrape won't have the DB row yet)
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir
        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        db = _DB(db_path)

        enriched: list[dict] = []
        for j in jobs:
            row = db.get_application(j.get("id")) if j.get("id") else None
            if row is None:
                row = dict(j)
            else:
                # Make sure job_url survives even if the row doesn't have it yet
                if not row.get("job_url") and j.get("job_url"):
                    row["job_url"] = j["job_url"]
            # Pull JD (cached if present, scraped and persisted otherwise)
            try:
                self._get_or_fetch_jd(row)
            except Exception as e:
                log.warning(
                    f"[JobSearch] JD fetch failed for #{row.get('id')}: {e}"
                )
            enriched.append(row)

        # Build prompt blocks — include JD excerpts for each job so Claude
        # can actually reason about requirements, not just title keywords.
        blocks: list[str] = []
        for j in enriched:
            lines = [
                f"--- tracker_id={j.get('id')} ---",
                f"company: {j.get('company', 'Unknown')}",
                f"position: {j.get('position', '')}",
            ]
            if j.get("location"):
                lines.append(f"location: {j['location']}")
            jd = (j.get("job_description") or "").strip()
            if jd:
                # Trim hard for multi-job batches; prefer start + end of JD
                if len(jd) > 3500:
                    jd = jd[:2500] + "\n[...]\n" + jd[-800:]
                lines.append(f"job_description:\n{jd}")
            else:
                lines.append("job_description: (not available)")
            blocks.append("\n".join(lines))

        prompt = (
            "TASK: Score each job listing for fit against this resume.\n\n"
            f"RESUME:\n{resume_text}\n\n"
            "JOB LISTINGS:\n\n"
            + "\n\n".join(blocks) + "\n\n"
            "Score fit 0-100 using the job_description text when available. "
            "Weight: seniority match (CISO / VP / Director level), "
            "security domain depth, regulated-industry experience, and "
            "stated required certifications (CISSP, CISM, etc.). "
            "80+ = strong match. 50-79 = partial. Below 50 = poor fit.\n\n"
            "Return ONLY a JSON array, nothing else. Each element:\n"
            '{"tracker_id": <int>, "fit_score": <int>, '
            '"analysis": "<2-3 sentences grounded in the JD>", '
            '"recommendation": "<apply|skip|maybe>"}'
        )

        try:
            claude_bin = shutil.which("claude")
            if not claude_bin:
                log.error("[JobSearch] claude CLI not found in PATH")
                return

            result = subprocess.run(
                [claude_bin, "-p", "--output-format", "text"],
                input=prompt,
                capture_output=True,
                text=True,
                encoding="utf-8",
                timeout=300,
                cwd=str(Path.home()),
            )

            if result.returncode != 0:
                log.error(
                    f"[JobSearch] claude -p failed (rc={result.returncode}): "
                    f"{result.stderr[:200]}"
                )
                return

            raw = result.stdout.strip()
            if not raw:
                log.error(
                    f"[JobSearch] claude -p returned empty output. "
                    f"stderr: {result.stderr[:300]}"
                )
                return

            # Strip markdown fences if Claude wraps them
            raw = re.sub(r"^```\w*\n?", "", raw)
            raw = re.sub(r"\n?```$", "", raw.strip())

            # Extract JSON array from anywhere in the response
            # (Claude sometimes adds preamble text before the JSON)
            arr_match = re.search(r'\[.*\]', raw, re.DOTALL)
            if not arr_match:
                log.error(
                    f"[JobSearch] No JSON array found in response. "
                    f"First 500 chars: {raw[:500]}"
                )
                return

            scores = json.loads(arr_match.group())
            if not isinstance(scores, list):
                log.error("[JobSearch] Fit analysis returned non-list JSON")
                return

            self._store_fit_results(scores)

        except subprocess.TimeoutExpired:
            log.error("[JobSearch] claude -p timed out after 300s")
        except json.JSONDecodeError as e:
            log.error(
                f"[JobSearch] Fit analysis JSON parse error: {e}. "
                f"First 500 chars: {raw[:500]}"
            )
        except Exception as e:
            log.error(f"[JobSearch] Fit analysis failed: {e}")

    @staticmethod
    def _store_fit_results(scores: list[dict]) -> None:
        """Write fit analysis results back to the job tracker database."""
        from talents.job_tracker import _DB, _data_dir as tracker_data_dir

        db_path = os.path.join(tracker_data_dir(), "job_tracker.db")
        db = _DB(db_path)

        updated = 0
        for entry in scores:
            app_id = entry.get("tracker_id")
            if not app_id:
                continue
            try:
                fields = {}
                if "fit_score" in entry:
                    fields["fit_score"] = int(entry["fit_score"])
                note_parts = []
                if entry.get("analysis"):
                    note_parts.append(entry["analysis"])
                if entry.get("recommendation"):
                    note_parts.append(
                        f"Recommendation: {entry['recommendation']}"
                    )
                if note_parts:
                    fields["notes"] = " | ".join(note_parts)
                if fields:
                    db.update_application(app_id, **fields)
                    updated += 1
            except Exception as e:
                log.error(
                    f"[JobSearch] Failed to store fit for #{app_id}: {e}"
                )

        log.info(f"[JobSearch] Fit analysis complete: {updated}/{len(scores)} updated")


# ── Filename / slug helpers ──────────────────────────────────────────────────

def _safe_slug(text: str, *, max_len: int = 50) -> str:
    """Turn a company or job title into a filesystem-safe slug.

    - HTML-unescapes (so '&amp;' becomes '&' before stripping)
    - Drops punctuation except word chars, whitespace, and hyphen
    - Collapses whitespace to single underscores
    - Truncates at a word boundary under max_len
    """
    import html as _html
    if not text:
        return "untitled"
    s = _html.unescape(str(text))
    s = re.sub(r'[^\w\s-]', '', s).strip()
    s = re.sub(r'\s+', '_', s)
    if len(s) <= max_len:
        return s or "untitled"
    cut = s[:max_len]
    # Truncate at last underscore before the cap
    last_us = cut.rfind('_')
    if last_us >= max_len // 2:
        cut = cut[:last_us]
    return cut.rstrip('_-') or "untitled"


# ── Response helpers ─────────────────────────────────────────────────────────

def _ok(response: str, *, actions: list | None = None) -> dict:
    return {
        "success": True,
        "response": response,
        "actions_taken": actions or [],
        "spoken": False,
    }


def _fail(response: str) -> dict:
    return {
        "success": False,
        "response": response,
        "actions_taken": [],
        "spoken": False,
    }
